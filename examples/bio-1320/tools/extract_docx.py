#!/usr/bin/env python3
"""Extract text from .docx files into markdown.

Usage:
  python3 tools/extract_docx.py [paths...]
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
EXTRACTED = REPO / "raw" / "extracted"


def slug_for(name: str) -> str:
    base = re.sub(r"\.docx$", "", name, flags=re.I)
    base = re.sub(r"\(\d+\)", "", base).strip()
    if "study guide" in base.lower() and "exam 3" in base.lower():
        return "study-guide-exam-3"
    return re.sub(r"[^a-z0-9-]+", "-", base.lower()).strip("-")


def is_heading(p) -> bool:
    return (p.style.name or "").lower().startswith("heading")


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    out: list[str] = [f"# {path.name}\n"]
    qcount = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            out.append("")
            continue
        # Headings
        if is_heading(p):
            level = 2
            try:
                level = int((p.style.name or "Heading 1").split()[-1]) + 1
            except Exception:
                pass
            level = min(max(level, 2), 4)
            out.append(f"{'#' * level} {text}\n")
            continue
        # Numbered question detection — give each a stable anchor
        m = re.match(r"^(\d+)[.)]\s+(.+)$", text)
        if m:
            qcount += 1
            num = m.group(1)
            rest = m.group(2)
            out.append(f"### Q{num} {{#q{num}}}\n")
            out.append(rest)
            out.append("")
            continue
        # Bullets
        if (p.style.name or "").lower().startswith("list"):
            out.append(f"- {text}")
            continue
        out.append(text)
    # Tables
    for ti, table in enumerate(doc.tables, 1):
        out.append("")
        out.append(f"### Table {ti}\n")
        rows = []
        for row in table.rows:
            rows.append([c.text.strip() for c in row.cells])
        if rows:
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for r in rows[1:]:
                out.append("| " + " | ".join(c.replace("\n", " ").replace("|", "\\|") for c in r) + " |")
            out.append("")
    return "\n".join(out)


def main() -> int:
    paths = [Path(p) for p in sys.argv[1:]]
    if not paths:
        for d in [REPO, REPO / "raw"]:
            if d.exists():
                paths.extend(sorted(d.glob("*.docx")))
        paths = list({p.resolve() for p in paths})

    if not paths:
        print("no .docx files found", file=sys.stderr)
        return 1

    EXTRACTED.mkdir(parents=True, exist_ok=True)
    for p in paths:
        slug = slug_for(p.name)
        out_path = EXTRACTED / f"{slug}.md"
        print(f"  {p.name} → {out_path.relative_to(REPO)}")
        out_path.write_text(extract_docx(p), encoding="utf-8")
    print(f"extracted {len(paths)} docx files into {EXTRACTED.relative_to(REPO)}/")
    return 0


if __name__ == "__main__":
    sys.exit(main())
